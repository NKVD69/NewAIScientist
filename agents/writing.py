"""
agents/writing.py — WritingAgent for scientific manuscript generation.

Responsible for:
- Drafting sections (Abstract, Introduction, Methods, Results, Discussion, Conclusion)
- Coordinating citations and bibliography
- Exporting to LaTeX and DOCX
"""

from __future__ import annotations

import logging
from typing import Any

from models.hypothesis import (
    Hypothesis,
    Manuscript,
    ManuscriptSection,
)
from utils.llm import get_llm_completion

from .base import BaseAgent

logger = logging.getLogger(__name__)


class WritingAgent(BaseAgent):
    """Drafts and compiles scientific manuscripts."""

    name = "Writing"

    async def draft_section(
        self,
        section_type: str,
        goal: Any,
        hypothesis: Hypothesis | None = None,
        context: dict | None = None
    ) -> ManuscriptSection:
        """Drafts a specific section of the paper."""
        if not self.llm_client:
            return ManuscriptSection(section_type=section_type, content=f"Drafting {section_type} requires LLM.")

        prompt = f"""You are a scientific writer for a top-tier journal (Nature/Science).
Section to write: {section_type}
Research Goal: {goal.title if hasattr(goal, 'title') else str(goal)}
Hypothesis: {hypothesis.title if hypothesis else 'N/A'}
Context/Data: {context if context else 'N/A'}

Write a rigorous, professional scientific {section_type}.
Use academic tone, citations like (Author et al., Year) if citations are provided in context.
Return ONLY the text content of the section."""

        try:
            response = await get_llm_completion(
                self.llm_client,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7,
                json_mode=False,
            )
            content = response.choices[0].message.content
            return ManuscriptSection(
                section_type=section_type,
                title=section_type.capitalize(),
                content=content
            )
        except Exception as e:
            logger.error(f"Drafting {section_type} failed: {e}")
            return ManuscriptSection(section_type=section_type, content=f"Error: {e}")

    async def compile_manuscript(
        self,
        goal: Any,
        sections: dict[str, ManuscriptSection],
        references: list[dict[str, str]]
    ) -> Manuscript:
        """Assembles all sections into a Manuscript model."""
        manuscript = Manuscript(
            title=goal.title if hasattr(goal, 'title') else "Scientific Paper",
            sections=sections,
            references=references
        )
        return manuscript

    # ------------------------------------------------------------------
    # EXPORTS
    # ------------------------------------------------------------------

    def export_to_latex(self, manuscript: Manuscript, filename: str = "paper.tex") -> str:
        """Generates a LaTeX file from the manuscript."""
        latex = r"\documentclass{article}" + "\n"
        latex += r"\usepackage[utf8]{inputenc}" + "\n"
        latex += r"\title{" + manuscript.title + "}\n"
        latex += r"\author{AI Co-Scientist Agent}\n"
        latex += r"\date{\today}\n"
        latex += r"\begin{document}" + "\n\n"
        latex += r"\maketitle" + "\n\n"

        # Order of sections
        order = ["abstract", "introduction", "methods", "results", "discussion", "conclusion"]
        for sec_name in order:
            if sec_name in manuscript.sections:
                sec = manuscript.sections[sec_name]
                if sec_name == "abstract":
                    latex += r"\begin{abstract}" + "\n"
                    latex += sec.content + "\n"
                    latex += r"\end{abstract}" + "\n\n"
                else:
                    latex += r"\section{" + sec.title + "}\n"
                    latex += sec.content + "\n\n"

        latex += r"\begin{thebibliography}{99}" + "\n"
        for i, ref in enumerate(manuscript.references, 1):
            title = ref.get("title", "Unknown Title")
            authors = ref.get("authors", "Unknown Authors")
            year = ref.get("year", "N/A")
            latex += fr"\bibitem{{ref{i}}} {authors} ({year}). {title}." + "\n"
        latex += r"\end{thebibliography}" + "\n\n"

        latex += r"\end{document}"

        with open(filename, "w", encoding="utf-8") as f:
            f.write(latex)

        return filename

    def export_to_docx(self, manuscript: Manuscript, filename: str = "paper.docx") -> str:
        """Generates a DOCX file from the manuscript using python-docx."""
        try:
            import docx
        except ImportError:
            logger.warning("python-docx not installed. Creating a simple text fallback.")
            with open(filename.replace(".docx", ".txt"), "w", encoding="utf-8") as f:
                f.write(f"TITLE: {manuscript.title}\n\n")
                for sec in manuscript.sections.values():
                    f.write(f"--- {sec.title} ---\n{sec.content}\n\n")
            return filename.replace(".docx", ".txt")

        doc = docx.Document()
        doc.add_heading(manuscript.title, 0)

        order = ["abstract", "introduction", "methods", "results", "discussion", "conclusion"]
        for sec_name in order:
            if sec_name in manuscript.sections:
                sec = manuscript.sections[sec_name]
                doc.add_heading(sec.title, level=1)
                doc.add_paragraph(sec.content)

        doc.add_heading("References", level=1)
        for i, ref in enumerate(manuscript.references, 1):
            title = ref.get("title", "Unknown Title")
            authors = ref.get("authors", "Unknown Authors")
            year = ref.get("year", "N/A")
            doc.add_paragraph(f"[{i}] {authors} ({year}). {title}.", style='List Bullet')

        doc.save(filename)
        return filename


__all__ = ["WritingAgent"]
